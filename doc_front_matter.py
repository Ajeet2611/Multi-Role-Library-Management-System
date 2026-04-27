"""
Front matter: Cover page, certificate, acknowledgement, abstract,
table of contents, list of figures, list of tables.
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from doc_helpers import *


def add_cover_page(doc):
    # Top decorative bar
    p = doc.add_paragraph()
    add_horizontal_line(p, "B8860B")
    p.paragraph_format.space_after = Pt(20)

    # Project Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROJECT DOCUMENTATION REPORT")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = GOLD
    p.paragraph_format.space_after = Pt(6)

    # Subtitle accent
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Technical & Architectural Reference Manual")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = MID_GREY
    p.paragraph_format.space_after = Pt(40)

    # Main project title in big
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Multi-Role")
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Library Management System")
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(20)

    # Hindi tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("एक Production-Grade Desktop Application")
    run.font.name = "Nirmala UI"
    run.font.size = Pt(18)
    run.font.italic = True
    run.font.color.rgb = GOLD
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Built with Python (Tkinter) + MySQL")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = MID_GREY
    p.paragraph_format.space_after = Pt(40)

    # Centered divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("◆◆◆")
    run.font.size = Pt(16)
    run.font.color.rgb = GOLD
    p.paragraph_format.space_after = Pt(40)

    # Author block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("प्रस्तुतकर्ता | Submitted By")
    run.font.name = "Nirmala UI"
    run.font.size = Pt(11)
    run.font.color.rgb = MID_GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ajeet Prasad")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Senior Developer | Full-Stack Engineer")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = MID_GREY
    p.paragraph_format.space_after = Pt(50)

    # Year/version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  •  April 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = MID_GREY

    # Bottom bar
    p = doc.add_paragraph()
    add_horizontal_line(p, "B8860B")
    p.paragraph_format.space_before = Pt(10)

    add_page_break(doc)


def add_certificate(doc):
    add_section_heading(doc, "", "प्रमाणपत्र | Certificate", level=1)
    div = doc.add_paragraph()
    add_horizontal_line(div, "B8860B")
    div.paragraph_format.space_after = Pt(20)

    add_hindi_para(doc,
        "यह प्रमाणित किया जाता है कि यह project report — \"Multi-Role Library "
        "Management System\" — Ajeet Prasad द्वारा तैयार की गई है। यह report एक "
        "comprehensive technical documentation है जो Python (Tkinter) और MySQL "
        "आधारित desktop application के design, implementation, और deployment "
        "के सभी पहलुओं को विस्तार से cover करती है।",
        size=12, first_line_indent=Inches(0.3),
        space_after=Pt(12))

    add_hindi_para(doc,
        "इस report में project की architecture, database design, security "
        "implementation, testing methodologies, और future scalability पर "
        "depth से discussion किया गया है। यह कार्य developer की technical "
        "expertise और engineering maturity का प्रमाण है।",
        size=12, first_line_indent=Inches(0.3),
        space_after=Pt(40))

    # Signature lines
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("\n\n_____________________                    _____________________")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GREY

    p = doc.add_paragraph()
    run = p.add_run("Project Guide                                              Head of Department")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GREY

    add_page_break(doc)


def add_acknowledgement(doc):
    add_section_heading(doc, "", "आभार | Acknowledgement", level=1)
    div = doc.add_paragraph()
    add_horizontal_line(div, "B8860B")
    div.paragraph_format.space_after = Pt(20)

    add_hindi_para(doc,
        "किसी भी project का successful completion अकेले possible नहीं होता — इसके "
        "पीछे कई लोगों का योगदान, मार्गदर्शन और support होता है। मैं इस अवसर पर "
        "उन सभी का धन्यवाद करना चाहता हूँ जिन्होंने इस project को possible बनाया।",
        size=12, first_line_indent=Inches(0.3))

    add_hindi_para(doc,
        "सबसे पहले मैं अपने Project Guide का धन्यवाद करता हूँ, जिनकी expert "
        "guidance और continuous feedback के बिना यह project complete नहीं हो "
        "सकता था। उनकी technical insights और experience ने मुझे सही direction "
        "में सोचने और काम करने की प्रेरणा दी।",
        first_line_indent=Inches(0.3))

    add_hindi_para(doc,
        "मैं अपने Head of Department और सभी faculty members का भी आभारी हूँ, "
        "जिन्होंने मुझे इस project को pursue करने का अवसर और infrastructure दिया।",
        first_line_indent=Inches(0.3))

    add_hindi_para(doc,
        "Open-source community का भी विशेष धन्यवाद — Python, MySQL, bcrypt, "
        "Tkinter, Matplotlib जैसे amazing tools जो इस project की foundation हैं, "
        "वो सब इन्हीं developers की मेहनत का result हैं।",
        first_line_indent=Inches(0.3))

    add_hindi_para(doc,
        "अंत में, मैं अपने family और friends का धन्यवाद करता हूँ जिनके support "
        "और patience के बिना यह 16-week journey complete नहीं हो सकती थी। उनकी "
        "encouragement मेरी सबसे बड़ी motivation रही है।",
        first_line_indent=Inches(0.3))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("\n— Ajeet Prasad")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = NAVY

    add_page_break(doc)


def add_abstract(doc):
    add_section_heading(doc, "", "सारांश | Abstract", level=1)
    div = doc.add_paragraph()
    add_horizontal_line(div, "B8860B")
    div.paragraph_format.space_after = Pt(20)

    add_hindi_para(doc,
        "Library Management आज भी भारत में एक significant challenge है — 60% से "
        "अधिक educational libraries अभी भी manual register-based systems पर "
        "depend करती हैं। यह project — Multi-Role Library Management System — "
        "इसी समस्या का एक comprehensive technical solution है।",
        size=12, first_line_indent=Inches(0.3))

    add_hindi_para(doc,
        "यह एक Python-based desktop application है जो Tkinter GUI framework और "
        "MySQL relational database का use करते हुए, libraries के सभी major "
        "operations को digitize करती है — book management, user authentication, "
        "issue/return tracking, automatic fine calculation, real-time analytics, "
        "और inter-user messaging।",
        first_line_indent=Inches(0.3))

    add_hindi_para(doc,
        "Project की key technical features में 3-Layer Architecture, bcrypt-based "
        "password hashing, Role-Based Access Control (RBAC), multi-institution "
        "isolation, और Matplotlib-based interactive dashboards शामिल हैं। "
        "Application 7+ MySQL tables पर operate करती है — सभी 3rd Normal Form "
        "में normalized — और 100% parameterized queries के through SQL injection "
        "से fully protected है।",
        first_line_indent=Inches(0.3))

    add_hindi_para(doc,
        "इस documentation में 16 chapters हैं जो project के हर pehlu को detail "
        "में explain करते हैं — initial problem statement से लेकर final deployment, "
        "security architecture से लेकर future roadmap तक। यह report एक production-"
        "grade software के design और development की complete journey को capture "
        "करती है।",
        first_line_indent=Inches(0.3))

    add_hindi_para(doc, "Keywords (English):", bold=True, color=NAVY,
                   size=11, space_after=Pt(2))
    add_para(doc, "Library Management System, Python, Tkinter, MySQL, bcrypt, "
             "Role-Based Access Control, Multi-Tenancy, 3-Layer Architecture, "
             "OWASP, Matplotlib, Desktop Application, Production-Grade Software",
             size=11, italic=True, color=MID_GREY)

    add_page_break(doc)


def add_table_of_contents(doc):
    add_section_heading(doc, "", "विषय सूची | Table of Contents", level=1)
    div = doc.add_paragraph()
    add_horizontal_line(div, "B8860B")
    div.paragraph_format.space_after = Pt(15)

    toc = [
        ("Front Matter", "", True),
        ("  Cover Page", "i", False),
        ("  Certificate", "ii", False),
        ("  Acknowledgement", "iii", False),
        ("  Abstract", "iv", False),
        ("  Table of Contents", "v", False),
        ("  List of Figures", "vii", False),
        ("  List of Tables", "viii", False),
        ("", "", False),
        ("Chapter 1 — परिचय (Introduction)", "1", True),
        ("  1.1  Project Overview", "1", False),
        ("  1.2  Background और Motivation", "2", False),
        ("  1.3  Project का Scope", "3", False),
        ("  1.4  Documentation का Structure", "4", False),
        ("  1.5  इस Document को कैसे पढ़ें", "5", False),
        ("", "", False),
        ("Chapter 2 — साहित्य समीक्षा (Literature Review)", "6", True),
        ("  2.1  Existing Systems Overview", "6", False),
        ("  2.2  Commercial Solutions", "7", False),
        ("  2.3  Comparative Analysis", "9", False),
        ("  2.4  Research Papers", "10", False),
        ("  2.5  Gap Analysis", "11", False),
        ("", "", False),
        ("Chapter 3 — समस्या कथन (Problem Statement)", "13", True),
        ("  3.1  Core Problem", "13", False),
        ("  3.2  Project Objectives", "14", False),
        ("  3.3  Success Criteria", "16", False),
        ("  3.4  Project Constraints", "17", False),
        ("", "", False),
        ("Chapter 4 — System Requirements", "18", True),
        ("  4.1  Hardware Requirements", "18", False),
        ("  4.2  Software Requirements", "19", False),
        ("  4.3  Functional Requirements", "20", False),
        ("  4.4  Non-Functional Requirements", "23", False),
        ("  4.5  User Roles and Permissions", "24", False),
        ("", "", False),
        ("Chapter 5 — Feasibility Study", "26", True),
        ("  5.1  Technical Feasibility", "26", False),
        ("  5.2  Economic Feasibility", "27", False),
        ("  5.3  Operational Feasibility", "29", False),
        ("  5.4  Schedule Feasibility", "30", False),
        ("  5.5  Risk Assessment", "31", False),
        ("", "", False),
        ("Chapter 6 — System Architecture", "33", True),
        ("  6.1  Architectural Overview", "33", False),
        ("  6.2  Presentation Layer", "34", False),
        ("  6.3  Business Logic Layer", "35", False),
        ("  6.4  Data Access Layer", "36", False),
        ("  6.5  Data Flow Example", "37", False),
        ("  6.6  Design Patterns", "38", False),
        ("", "", False),
        ("Chapter 7 — Tech Stack Deep Dive", "40", True),
        ("  7.1  Selection Philosophy", "40", False),
        ("  7.2  Python 3.10+", "40", False),
        ("  7.3  Tkinter", "41", False),
        ("  7.4  MySQL", "42", False),
        ("  7.5  bcrypt", "43", False),
        ("  7.6  Matplotlib", "44", False),
        ("", "", False),
        ("Chapter 8 — Database Design", "46", True),
        ("  8.1  Design Philosophy", "46", False),
        ("  8.2  ER Diagram", "47", False),
        ("  8.3  Table Specifications", "48", False),
        ("  8.4  Normalization (3NF)", "52", False),
        ("  8.5  Indexes", "53", False),
        ("  8.6  Multi-Tenancy", "54", False),
        ("", "", False),
        ("Chapter 9 — Module Implementation", "56", True),
        ("  9.1  Project Structure", "56", False),
        ("  9.2  db.py Walkthrough", "57", False),
        ("  9.3  security.py Walkthrough", "58", False),
        ("  9.4  auth.py Walkthrough", "59", False),
        ("  9.5  permissions.py Walkthrough", "60", False),
        ("  9.6  messaging.py Walkthrough", "61", False),
        ("  9.7  dashboard_charts.py", "62", False),
        ("  9.8-9.10  Admin Modules", "63", False),
        ("", "", False),
        ("Chapter 10 — Security Implementation", "66", True),
        ("  10.1  Security First Approach", "66", False),
        ("  10.2  bcrypt Layer", "66", False),
        ("  10.3  RBAC Layer", "67", False),
        ("  10.4  SQL Injection Prevention", "68", False),
        ("  10.5  Multi-Tenant Isolation", "69", False),
        ("  10.6  OWASP Compliance", "70", False),
        ("", "", False),
        ("Chapter 11 — UI/UX Design", "72", True),
        ("  11.1  Design Principles", "72", False),
        ("  11.2  Color Palette", "72", False),
        ("  11.3-11.12  Screen Walkthroughs", "73", False),
        ("", "", False),
        ("Chapter 12 — Testing & Validation", "78", True),
        ("  12.1  Testing Strategy", "78", False),
        ("  12.2  Unit Testing", "78", False),
        ("  12.3  Integration Testing", "79", False),
        ("  12.4  System Testing", "80", False),
        ("  12.5  UAT", "81", False),
        ("  12.6  Test Coverage", "82", False),
        ("", "", False),
        ("Chapter 13 — Results & Output", "84", True),
        ("  13.1-13.7  Application Screenshots", "84", False),
        ("  13.8  Performance Metrics", "92", False),
        ("", "", False),
        ("Chapter 14 — Challenges & Limitations", "94", True),
        ("  14.1  Engineering Challenges", "94", False),
        ("  14.2  Current Limitations", "97", False),
        ("  14.3  Lessons Learned", "98", False),
        ("", "", False),
        ("Chapter 15 — Future Enhancements", "99", True),
        ("  15.1-15.4  Roadmap", "99", False),
        ("  15.5  Architecture Evolution", "102", False),
        ("", "", False),
        ("Chapter 16 — Conclusion", "104", True),
        ("", "", False),
        ("References & Appendix", "107", True),
    ]

    for entry, page, is_chapter in toc:
        if not entry:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        # Title
        run = p.add_run(entry)
        run.font.name = "Calibri" if is_chapter else "Calibri"
        run.font.size = Pt(11 if is_chapter else 10)
        run.font.bold = is_chapter
        run.font.color.rgb = NAVY if is_chapter else DARK_GREY
        # Dots and page
        if page:
            dots = p.add_run("  " + "." * max(2, 70 - len(entry)) + "  ")
            dots.font.name = "Calibri"
            dots.font.size = Pt(9)
            dots.font.color.rgb = MID_GREY
            page_run = p.add_run(page)
            page_run.font.name = "Calibri"
            page_run.font.size = Pt(10)
            page_run.font.bold = is_chapter
            page_run.font.color.rgb = NAVY if is_chapter else DARK_GREY

    add_page_break(doc)


def add_list_of_figures(doc):
    add_section_heading(doc, "", "Figure सूची | List of Figures", level=1)
    div = doc.add_paragraph()
    add_horizontal_line(div, "B8860B")
    div.paragraph_format.space_after = Pt(15)

    figures = [
        ("Figure 3.1", "Problem-Solution Mapping Diagram", "17"),
        ("Figure 6.1", "3-Layer Architecture Diagram", "34"),
        ("Figure 6.2", "Login Process Data Flow", "37"),
        ("Figure 6.3", "Application Architecture Visualization", "39"),
        ("Figure 7.1", "Tech Stack Visualization", "45"),
        ("Figure 8.1", "Complete ER Diagram", "47"),
        ("Figure 8.2", "MySQL Workbench Schema", "55"),
        ("Figure 9.1", "Complete Project Structure", "56"),
        ("Figure 9.2", "VS Code Project File Tree", "65"),
        ("Figure 10.1", "Security Architecture Diagram", "71"),
        ("Figure 11.1", "Login Window Screenshot", "73"),
        ("Figure 11.2", "Super Admin Dashboard", "74"),
        ("Figure 11.3", "Admin Dashboard - Books", "74"),
        ("Figure 11.4", "Admin - Issue Book Form", "75"),
        ("Figure 11.5", "Admin - Return Book", "75"),
        ("Figure 11.6", "User Dashboard", "75"),
        ("Figure 11.7", "Add Book Form", "76"),
        ("Figure 11.8", "Books List View", "76"),
        ("Figure 11.9", "User Messaging Window", "76"),
        ("Figure 11.10", "Admin Inbox", "76"),
        ("Figure 11.11", "Pie Chart Analytics", "77"),
        ("Figure 11.12", "Bar Chart Analytics", "77"),
        ("Figure 11.13", "User Activity Log", "77"),
        ("Figure 11.14", "Permission Manager", "77"),
        ("Figure 12.1", "Test Execution Console", "83"),
        ("Figure 13.1-13.25", "Application Output Screenshots", "84-91"),
        ("Figure 15.1", "Future Roadmap Visualization", "103"),
    ]
    for fig_id, caption, page in figures:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{fig_id}  ")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run = p.add_run(caption)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
        dots = p.add_run("  " + "." * max(2, 60 - len(caption)) + "  ")
        dots.font.size = Pt(8)
        dots.font.color.rgb = MID_GREY
        run = p.add_run(page)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
    add_page_break(doc)


def add_list_of_tables(doc):
    add_section_heading(doc, "", "Table सूची | List of Tables", level=1)
    div = doc.add_paragraph()
    add_horizontal_line(div, "B8860B")
    div.paragraph_format.space_after = Pt(15)

    tables = [
        ("Table 1", "Project Scope - Functional Areas", "3"),
        ("Table 2", "Documentation Structure", "4"),
        ("Table 3", "Comparative Analysis vs Existing Solutions", "10"),
        ("Table 4", "Project Success Criteria", "16"),
        ("Table 5", "Hardware Requirements", "18"),
        ("Table 6", "Software Requirements", "19"),
        ("Table 7", "Non-Functional Requirements", "23"),
        ("Table 8", "User Roles & Permissions Matrix", "25"),
        ("Table 9", "Project Cost Breakdown", "28"),
        ("Table 10", "Project Schedule (16-week)", "31"),
        ("Table 11", "Risk Assessment Matrix", "32"),
        ("Table 12", "Presentation Layer Files", "35"),
        ("Table 13", "Design Patterns Implemented", "38"),
        ("Table 14", "GUI Framework Comparison", "42"),
        ("Table 15", "Database Comparison", "43"),
        ("Table 16", "Password Hashing Comparison", "44"),
        ("Table 17", "Supporting Python Libraries", "45"),
        ("Tables 18-23", "Database Table Schemas", "48-52"),
        ("Table 24", "Role Hierarchy", "67"),
        ("Table 25", "OWASP Top 10 Mitigation", "70"),
        ("Table 26", "UI Color Palette", "72"),
        ("Tables 27-31", "Test Cases & Bug Tracking", "78-83"),
        ("Table 32", "Performance Benchmark", "92"),
        ("Tables 33-35", "Future Roadmap", "99-101"),
        ("Table 36", "Skills Demonstrated", "105"),
        ("Tables 37-39", "Appendix Tables", "108-110"),
    ]
    for tbl_id, caption, page in tables:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{tbl_id}  ")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run = p.add_run(caption)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
        dots = p.add_run("  " + "." * max(2, 60 - len(caption)) + "  ")
        dots.font.size = Pt(8)
        dots.font.color.rgb = MID_GREY
        run = p.add_run(page)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
    add_page_break(doc)
