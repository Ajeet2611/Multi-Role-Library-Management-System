"""
Helper functions for building the Library Management System documentation.
Provides reusable styling, headings, code blocks, screenshot placeholders, etc.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION


# ===== Theme colors =====
NAVY = RGBColor(0x0B, 0x1F, 0x3A)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
MID_GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY_BG = "F2F2F2"
CODE_BG = "1E2A3A"
CODE_FG = RGBColor(0xE8, 0xE8, 0xE8)
PLACEHOLDER_BG = "EEF3FA"
PLACEHOLDER_BORDER = "5A7BA8"
ACCENT_BLUE = RGBColor(0x1F, 0x4E, 0x79)


def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex="000000", size="6"):
    """Add borders to a single table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_horizontal_line(paragraph, color="B8860B"):
    """Insert a bottom border on a paragraph (acts as a horizontal divider)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_break(doc):
    doc.add_page_break()


def add_chapter_heading(doc, chapter_num, title_hindi, title_english=None):
    """Big chapter heading on its own with a horizontal divider."""
    # Chapter label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"CHAPTER {chapter_num}")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = GOLD
    run.font.all_caps = True

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title_hindi)
    run.font.name = "Mangal"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY

    if title_english:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"({title_english})")
        run2.font.name = "Calibri"
        run2.font.size = Pt(14)
        run2.font.italic = True
        run2.font.color.rgb = MID_GREY

    # Horizontal divider
    div = doc.add_paragraph()
    add_horizontal_line(div, "B8860B")
    div.paragraph_format.space_after = Pt(12)


def add_section_heading(doc, num, title, level=1):
    """Section/sub-section heading. Level 1 = X.Y, Level 2 = X.Y.Z."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        run = p.add_run(f"{num}  {title}")
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY
    elif level == 2:
        run = p.add_run(f"{num}  {title}")
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = ACCENT_BLUE
    else:
        run = p.add_run(f"{num}  {title}")
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = DARK_GREY


def add_para(doc, text, size=11, bold=False, italic=False, align=None,
             color=DARK_GREY, font="Calibri", space_after=Pt(8),
             first_line_indent=None):
    """Add a regular paragraph. Use 'Mangal' or 'Nirmala UI' font for Devanagari."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    run.font.name = font
    # Apply complex script font for Devanagari
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:cs"), "Nirmala UI")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_hindi_para(doc, text, size=11, bold=False, italic=False, align=None,
                   space_after=Pt(8), color=DARK_GREY, first_line_indent=None):
    """Optimized for Hindi (Devanagari) text — uses Nirmala UI fallback."""
    return add_para(doc, text, size=size, bold=bold, italic=italic, align=align,
                    color=color, font="Nirmala UI", space_after=space_after,
                    first_line_indent=first_line_indent)


def add_bullet(doc, text, level=0, size=11, font="Nirmala UI"):
    """Add a bullet point (uses Word's built-in List Bullet style)."""
    style_name = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:cs"), "Nirmala UI")
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GREY
    return p


def add_numbered(doc, text, size=11, font="Nirmala UI"):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:cs"), "Nirmala UI")
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GREY


def add_code_block(doc, code, language="python"):
    """Insert a styled code block (single-cell table with dark background)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, CODE_BG)
    set_cell_borders(cell, "000000", "4")
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Language label
    label_run = p.add_run(f"  {language.upper()}\n")
    label_run.font.name = "Consolas"
    label_run.font.size = Pt(8)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
    # Code
    code_run = p.add_run(code)
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = CODE_FG
    # spacing after
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_screenshot_placeholder(doc, figure_num, caption,
                               height_inches=3.5):
    """A bordered placeholder box where a screenshot can be inserted later."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    set_cell_bg(cell, PLACEHOLDER_BG)
    set_cell_borders(cell, PLACEHOLDER_BORDER, "12")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Force row height
    tr = cell._tc.getparent()
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_inches * 1440)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n").font.size = Pt(int(height_inches * 6))
    icon = p.add_run("📷\n\n")
    icon.font.size = Pt(36)
    icon.font.color.rgb = ACCENT_BLUE
    p.add_run("\n").font.size = Pt(6)
    label = p.add_run("[ Screenshot यहाँ Insert करें ]\n")
    label.font.name = "Calibri"
    label.font.size = Pt(11)
    label.font.bold = True
    label.font.color.rgb = ACCENT_BLUE
    sub = p.add_run(caption)
    sub.font.name = "Nirmala UI"
    sub.font.size = Pt(9.5)
    sub.font.italic = True
    sub.font.color.rgb = MID_GREY

    # Caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figure {figure_num}: {caption}")
    cap_run.font.name = "Calibri"
    cap_run.font.size = Pt(10)
    cap_run.font.bold = True
    cap_run.font.italic = True
    cap_run.font.color.rgb = NAVY
    cap.paragraph_format.space_after = Pt(12)


def add_diagram_placeholder(doc, figure_num, caption,
                            ascii_diagram, height_inches=4.0):
    """Diagram with optional ASCII art preview inside."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_bg(cell, "FAFAFA")
    set_cell_borders(cell, "B8860B", "12")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code_run = p.add_run(ascii_diagram)
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = NAVY

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figure {figure_num}: {caption}")
    cap_run.font.name = "Calibri"
    cap_run.font.size = Pt(10)
    cap_run.font.bold = True
    cap_run.font.italic = True
    cap_run.font.color.rgb = NAVY
    cap.paragraph_format.space_after = Pt(12)


def add_data_table(doc, headers, rows, table_num=None, caption=None,
                   col_widths=None, header_color="0B1F3A"):
    """Create a styled data table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        set_cell_borders(cell, "0B1F3A", "8")
        if col_widths and i < len(col_widths):
            cell.width = col_widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Body
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        zebra = LIGHT_GREY_BG if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, zebra)
            set_cell_borders(cell, "B0B0B0", "4")
            if col_widths and c_idx < len(col_widths):
                cell.width = col_widths[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Nirmala UI"
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GREY

    # Caption
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prefix = f"Table {table_num}: " if table_num else ""
        cap_run = cap.add_run(f"{prefix}{caption}")
        cap_run.font.name = "Calibri"
        cap_run.font.size = Pt(10)
        cap_run.font.bold = True
        cap_run.font.italic = True
        cap_run.font.color.rgb = NAVY
        cap.paragraph_format.space_after = Pt(12)
    else:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)


def add_callout_box(doc, title, content, color="FFF8E1", border="B8860B"):
    """Highlighted callout box (info, warning, pro-tip etc.)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_bg(cell, color)
    set_cell_borders(cell, border, "12")
    p = cell.paragraphs[0]
    title_run = p.add_run(f"💡 {title}\n")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(11)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY
    body_run = p.add_run(content)
    body_run.font.name = "Nirmala UI"
    body_run.font.size = Pt(10.5)
    body_run.font.color.rgb = DARK_GREY
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_page_number_footer(section):
    """Add page number to the section footer."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GREY


def add_header(section, text):
    """Add a header to the section."""
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MID_GREY
    add_horizontal_line(p, "B8860B")
