"""
Multi-Role Library Management System — Comprehensive Documentation Generator
Generates a 70-80 page professional Hindi/Hinglish documentation in .docx format.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from doc_helpers import (
    add_page_number_footer, add_header, NAVY, GOLD
)
from doc_front_matter import (
    add_cover_page, add_certificate, add_acknowledgement,
    add_abstract, add_table_of_contents,
    add_list_of_figures, add_list_of_tables,
)
from doc_chapters_1_5 import (
    chapter_1_introduction, chapter_2_literature_review,
    chapter_3_problem_statement, chapter_4_requirements,
    chapter_5_feasibility,
)
from doc_chapters_6_9 import (
    chapter_6_architecture, chapter_7_tech_stack,
    chapter_8_database, chapter_9_modules,
)
from doc_chapters_10_16 import (
    chapter_10_security, chapter_11_uiux, chapter_12_testing,
    chapter_13_results, chapter_14_challenges, chapter_15_future,
    chapter_16_conclusion, chapter_17_references_appendix,
)


def setup_document_style(doc):
    """Configure default styles, margins, page size."""
    # Page setup
    for section in doc.sections:
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.4
    style.paragraph_format.space_after = Pt(8)


def main():
    print("=" * 60)
    print("Library Management System - Documentation Generator")
    print("=" * 60)

    doc = Document()
    setup_document_style(doc)

    # Configure footer/header for all sections
    add_page_number_footer(doc.sections[0])
    add_header(doc.sections[0],
               "Multi-Role Library Management System | Documentation")

    print("\n[1/8] Building front matter...")
    add_cover_page(doc)
    add_certificate(doc)
    add_acknowledgement(doc)
    add_abstract(doc)
    add_table_of_contents(doc)
    add_list_of_figures(doc)
    add_list_of_tables(doc)

    print("[2/8] Chapter 1-2: Introduction & Literature Review...")
    chapter_1_introduction(doc)
    chapter_2_literature_review(doc)

    print("[3/8] Chapter 3-5: Problem, Requirements, Feasibility...")
    chapter_3_problem_statement(doc)
    chapter_4_requirements(doc)
    chapter_5_feasibility(doc)

    print("[4/8] Chapter 6-7: Architecture & Tech Stack...")
    chapter_6_architecture(doc)
    chapter_7_tech_stack(doc)

    print("[5/8] Chapter 8-9: Database & Modules...")
    chapter_8_database(doc)
    chapter_9_modules(doc)

    print("[6/8] Chapter 10-12: Security, UI/UX, Testing...")
    chapter_10_security(doc)
    chapter_11_uiux(doc)
    chapter_12_testing(doc)

    print("[7/8] Chapter 13-16: Results, Challenges, Future, Conclusion...")
    chapter_13_results(doc)
    chapter_14_challenges(doc)
    chapter_15_future(doc)
    chapter_16_conclusion(doc)

    print("[8/8] Chapter 17: References & Appendix...")
    chapter_17_references_appendix(doc)

    # Save
    os.makedirs("output", exist_ok=True)
    output_path = "output/Library_Management_System_Documentation.docx"
    doc.save(output_path)

    file_size_kb = os.path.getsize(output_path) / 1024
    print("\n" + "=" * 60)
    print(f"✓ Documentation generated successfully!")
    print(f"✓ File: {output_path}")
    print(f"✓ Size: {file_size_kb:.1f} KB")
    print(f"✓ Estimated pages: 90+")
    print("=" * 60)


if __name__ == "__main__":
    main()
